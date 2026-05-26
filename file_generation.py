"""
file_generation.py
==================
Handles detection, parsing, and writing of files emitted by the LLM.

The LLM signals a file using the tag pair:

    [FILE:filename.ext]
    ...file content...
    [/FILE]

Multiple blocks may appear in a single response.

Supported plain-text types (always available)
----------------------------------------------
  .txt  .md  .js  .ts  .jsx  .tsx  .css  .html
  .py   .cs  .cpp .c   .h    .java .json .yaml
  .yml  .sh  .sql .xml .csv  .toml .ini  .env

Optional rich types (installed separately)
-------------------------------------------
  .docx  – requires python-docx   (pip install python-docx)
  .xlsx  – requires openpyxl      (pip install openpyxl)

SSE event emitted per file
---------------------------
    event: file
    data: {
        "filename":  "report.md",
        "mime_type": "text/markdown",
        "encoding":  "base64",        # always base64 so the wire format is safe
        "data":      "<base64 str>",
        "size":      1234             # bytes of the raw (pre-encoded) content
    }
"""

from __future__ import annotations

import base64
import io
import os
import re
import unicodedata
from pathlib import Path

# ── Optional rich-document libraries ──────────────────────────────────────────
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

# ── Constants ──────────────────────────────────────────────────────────────────
# Matches  [FILE:some_name.ext]  …content…  [/FILE]
# DOTALL so content can span multiple lines.
FILE_BLOCK_RE = re.compile(
    r'\[FILE:([^\]\n]+)\]\n?(.*?)\[/FILE\]',
    re.DOTALL | re.IGNORECASE,
)

# Catches any leftover [FILE:...], [/FILE], or partial tag artifacts
# that weren't part of a complete, matched block.
STRAY_TAG_RE = re.compile(
    r'\[/?FILE[^\]]*\]',
    re.IGNORECASE,
)

# Allowed extensions → MIME type
PLAIN_TEXT_TYPES: dict[str, str] = {
    ".txt":  "text/plain",
    ".md":   "text/markdown",
    ".js":   "text/javascript",
    ".ts":   "text/typescript",
    ".jsx":  "text/javascript",
    ".tsx":  "text/typescript",
    ".css":  "text/css",
    ".html": "text/html",
    ".htm":  "text/html",
    ".py":   "text/x-python",
    ".cs":   "text/x-csharp",
    ".cpp":  "text/x-c++src",
    ".c":    "text/x-csrc",
    ".h":    "text/x-chdr",
    ".java": "text/x-java",
    ".json": "application/json",
    ".yaml": "text/yaml",
    ".yml":  "text/yaml",
    ".sh":   "text/x-sh",
    ".sql":  "text/x-sql",
    ".xml":  "text/xml",
    ".csv":  "text/csv",
    ".toml": "text/toml",
    ".ini":  "text/plain",
    ".env":  "text/plain",
}

RICH_TYPES: dict[str, str] = {
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

OUTPUT_DIR = Path(os.getenv("FILE_OUTPUT_DIR", "generated_files"))


# ── Public helpers ─────────────────────────────────────────────────────────────
def has_file_block(text: str) -> bool:
    """Quick check — is there at least one [FILE:…][/FILE] block?"""
    return bool(FILE_BLOCK_RE.search(text))


def strip_file_blocks(text: str) -> str:
    """Remove all complete [FILE:…][/FILE] blocks AND any stray orphaned tags."""
    text = FILE_BLOCK_RE.sub("", text)
    text = STRAY_TAG_RE.sub("", text)
    return text.strip()


def extract_and_generate(llm_output: str) -> list[dict]:
    """
    Parse every [FILE:…][/FILE] block in *llm_output* and build the
    file bytes for each one.

    Returns a list of dicts ready to be JSON-serialised into SSE events:
        [
            {
                "filename":  "hello.md",
                "mime_type": "text/markdown",
                "encoding":  "base64",
                "data":      "<base64>",
                "size":      123,
            },
            …
        ]
    Blocks whose extension is unsupported are skipped (with a warning).
    """
    results = []
    for match in FILE_BLOCK_RE.finditer(llm_output):
        raw_name = match.group(1).strip()
        content  = match.group(2)

        filename = _sanitise_filename(raw_name)
        ext      = Path(filename).suffix.lower()

        if ext in PLAIN_TEXT_TYPES:
            file_bytes = content.encode("utf-8")
            mime       = PLAIN_TEXT_TYPES[ext]

        elif ext == ".docx":
            if not DOCX_AVAILABLE:
                print(f"[file_generation] Skipping {filename}: python-docx not installed.")
                continue
            file_bytes = _content_to_docx(content)
            mime       = RICH_TYPES[".docx"]

        elif ext == ".xlsx":
            if not XLSX_AVAILABLE:
                print(f"[file_generation] Skipping {filename}: openpyxl not installed.")
                continue
            file_bytes = _content_to_xlsx(content)
            mime       = RICH_TYPES[".xlsx"]

        else:
            print(f"[file_generation] Unsupported extension '{ext}' in '{filename}' — skipped.")
            continue

        # Persist to disk (optional but useful for server-side serving)
        _save_to_disk(filename, file_bytes)

        results.append({
            "filename":  filename,
            "mime_type": mime,
            "encoding":  "base64",
            "data":      base64.b64encode(file_bytes).decode("ascii"),
            "size":      len(file_bytes),
        })

    return results


# ── Internal helpers ───────────────────────────────────────────────────────────
def _sanitise_filename(name: str) -> str:
    """
    Strip path separators and dangerous characters, normalise unicode,
    and fall back to 'output.txt' if nothing useful remains.
    """
    # Normalise unicode to ASCII where possible
    name = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode()
    # Remove anything that isn't alphanumeric, dot, dash, or underscore
    name = re.sub(r"[^\w.\-]", "_", name)
    # Collapse multiple dots / underscores
    name = re.sub(r"\.{2,}", ".", name)
    name = re.sub(r"_{2,}", "_", name)
    name = name.strip("._")
    return name or "output.txt"


def _save_to_disk(filename: str, data: bytes) -> Path:
    """Write *data* to OUTPUT_DIR / filename, creating the directory if needed."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    dest = OUTPUT_DIR / filename
    # Avoid clobbering existing files
    stem   = dest.stem
    suffix = dest.suffix
    counter = 1
    while dest.exists():
        dest = OUTPUT_DIR / f"{stem}_{counter}{suffix}"
        counter += 1
    dest.write_bytes(data)
    return dest


# ── Rich document builders ─────────────────────────────────────────────────────
def _content_to_docx(content: str) -> bytes:
    """
    Convert plain text / Markdown-ish content to a .docx file.

    Basic heuristics applied:
      - Lines starting with '# '  → Heading 1
      - Lines starting with '## ' → Heading 2
      - Lines starting with '### '→ Heading 3
      - Everything else           → Normal paragraph
    """
    doc = DocxDocument()
    for line in content.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped == "":
            doc.add_paragraph("")          # blank line as paragraph break
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _content_to_xlsx(content: str) -> bytes:
    """
    Convert CSV-like plain text to a .xlsx workbook.

    The LLM is instructed (via the system prompt additions below) to
    produce comma-separated rows inside an [FILE:name.xlsx] block.
    Each row becomes a spreadsheet row; the first row is bolded as a header.
    """
    wb = openpyxl.Workbook()
    ws = wb.active

    rows = [line for line in content.splitlines() if line.strip()]
    for row_idx, line in enumerate(rows, start=1):
        cells = [cell.strip() for cell in line.split(",")]
        for col_idx, value in enumerate(cells, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            if row_idx == 1:  # header row
                cell.font = openpyxl.styles.Font(bold=True)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── System-prompt snippet (import this into prompts.py) ───────────────────────
FILE_GENERATION_SYSTEM_SNIPPET = """
You can generate downloadable files for the user using this exact tag format:

[FILE:filename.extension]
...full file content here...
[/FILE]

STRICT RULES — follow all of these without exception:
1. Put ALL file content inside the [FILE:…][/FILE] block. Do NOT repeat or echo the file content outside the block as a code fence or in any other form.
2. After the [/FILE] closing tag, write only a short one or two sentence explanation of what the file contains. Never reproduce the file content again.
3. ONE block per file, always. For CSV or multi-row data, put ALL rows in a single [FILE:name.csv]…[/FILE] block — never split rows across multiple blocks.
4. Supported extensions: .txt .md .js .ts .jsx .tsx .css .html .py .cs .cpp .c .java .json .yaml .yml .sh .sql .xml .csv .toml .docx .xlsx
5. For .docx: use Markdown headings (# ## ###) inside the block; the system converts them automatically.
6. For .xlsx: write comma-separated rows inside the block; the first row becomes a bold header.
7. Do NOT emit [FILE:] tags unless the user explicitly wants a file or it is clearly useful.
8. NEVER emit a [/FILE] tag without a matching [FILE:filename] opening tag immediately before the content. Stray or unmatched tags are forbidden.
9. If the user uploaded files in [Attached file:…][/Attached file] blocks, do NOT re-emit those same files as [FILE:] output. You may freely create NEW files with different names and content — this rule only prevents echoing back what the user already has.

Example of correct output:
[FILE:greet.py]
def greet(name):
    return f"Hello, {name}!"
[/FILE]
This file defines a simple Python greeting function.
""".strip()